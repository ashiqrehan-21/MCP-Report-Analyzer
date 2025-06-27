from mcp.server.fastmcp import FastMCP
import os
import docx
import smtplib
from email.message import EmailMessage
from typing import Optional
from collections import Counter
import sys

# Create an MCP server with SSE transport enabled
mcp = FastMCP("Doc Analyzer", use_sse=True)

REPORTS_DIR = os.path.join(os.path.dirname(__file__), '..', 'reports')

def get_report_path(filename: str) -> str:
    file_path = os.path.join(REPORTS_DIR, filename)
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Report file '{filename}' not found in reports directory.")
    return file_path

@mcp.tool()
def read_text_report(filename: str) -> str:
    """
    Reads the content of a text document report.

    Args:
        filename: The name of the report file in the 'reports' directory.

    Returns:
        The content of the document.
    """
    try:
        file_path = get_report_path(filename)
        with open(file_path, 'r') as f:
            return f.read()
    except FileNotFoundError as e:
        return str(e)
    except Exception as e:
        return f"An error occurred: {e}"

@mcp.tool()
def read_docx_report(filename: str) -> str:
    """
    Reads the content of a .docx document report.

    Args:
        filename: The name of the .docx report file in the 'reports' directory.

    Returns:
        The text content of the document.
    """
    try:
        file_path = get_report_path(filename)
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    except FileNotFoundError as e:
        return str(e)
    except Exception as e:
        return f"An error occurred: {e}"

@mcp.tool()
def produce_report_summary(filename: str) -> str:
    """
    Reads a .docx report and produces a summary of its findings by
    looking for sections like 'Executive Summary' or 'Conclusion'.

    Args:
        filename: The name of the .docx report file in the 'reports' directory.

    Returns:
        A summary of the report's findings, or a preview if no summary section is found.
    """
    try:
        file_path = get_report_path(filename)
        doc = docx.Document(file_path)
        
        summary_sections = [
            "Executive Summary", 
            "Summary of Findings", 
            "Conclusion", 
            "Summary"
        ]
        
        # A simple state machine to find and capture text under a summary heading
        is_capturing = False
        summary_paragraphs = []
        
        for para in doc.paragraphs:
            # If we find a heading that matches one of our sections, start capturing
            if any(section.lower() in para.text.lower() for section in summary_sections):
                is_capturing = True
                continue # Skip the heading itself
            
            # If we're capturing, add the paragraph text
            if is_capturing:
                # Stop if we hit a new major heading or a long blank space,
                # suggesting the end of the section.
                if para.style.name.startswith('Heading') and summary_paragraphs:
                    break
                if para.text.strip(): # Only add non-empty paragraphs
                    summary_paragraphs.append(para.text)

        if summary_paragraphs:
            return "\n".join(summary_paragraphs)
        else:
            # Fallback: if no summary section was found, return the first 3 paragraphs as a preview
            preview = [p.text for p in doc.paragraphs if p.text.strip()][:3]
            return "No summary section found. Preview:\n" + "\n".join(preview)

    except FileNotFoundError as e:
        return str(e)
    except Exception as e:
        return f"An error occurred: {e}"

# Helper to extract defect counts and severity breakdown from the report
def extract_defect_insights(filename: str):
    try:
        file_path = get_report_path(filename)
        doc = docx.Document(file_path)
        # Vulnerability severities to look for
        severities = ["Critical", "Medium", "Low", "Informational"]
        # Map to count each severity
        severity_counts = {s: 0 for s in severities}
        total = 0
        for para in doc.paragraphs:
            for sev in severities:
                if sev.lower() in para.text.lower():
                    severity_counts[sev] += 1
                    total += 1
        return total, severity_counts
    except Exception as e:
        return 0, {s: 0 for s in ["Critical", "Medium", "Low", "Informational"]}

# Helper to extract high/critical/medium defect counts
def extract_high_critical_medium_counts(filename: str):
    try:
        file_path = get_report_path(filename)
        doc = docx.Document(file_path)
        # Vulnerability severities to look for
        severities = ["Critical", "High", "Medium"]
        counts = Counter({s: 0 for s in severities})
        for para in doc.paragraphs:
            for sev in severities:
                if sev.lower() in para.text.lower():
                    counts[sev] += 1
        return dict(counts)
    except Exception as e:
        return {s: 0 for s in ["Critical", "High", "Medium"]}

@mcp.tool()
def send_report_summary_email(recipient: str, filename: str, smtp_server: str = "smtp.gmail.com", smtp_port: int = 587, smtp_user: str = '', smtp_password: str = '', custom_message: Optional[str] = None) -> str:
    """
    Sends an email with the report summary and insights (defect count, severity breakdown).

    Args:
        recipient: Email address to send the report to.
        filename: The name of the .docx report file in the 'reports' directory.
        smtp_server: SMTP server address.
        smtp_port: SMTP server port.
        smtp_user: SMTP username (must be provided, not empty).
        smtp_password: SMTP password (must be provided, not empty).
        custom_message: Optional custom message to include in the email.

    Returns:
        Success or error message.
    """
    if not smtp_user or not smtp_password:
        return "SMTP username and password must be provided."
    summary = produce_report_summary(filename)
    total, severity_counts = extract_defect_insights(filename)
    high_crit_med = extract_high_critical_medium_counts(filename)
    insights = f"Total Defects: {total}\n" + "\n".join([f"{k}: {v}" for k, v in severity_counts.items()])
    body = f"Report Summary:\n{summary}\n\nInsights:\n{insights}"
    if custom_message is None:
        custom_message = ""
    if custom_message:
        body = f"{custom_message}\n\n" + body
    msg = EmailMessage()
    msg["Subject"] = f"Penetration Test Report Summary: {filename}"
    msg["From"] = smtp_user
    msg["To"] = recipient
    msg.set_content(body)
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
        return f"Email sent to {recipient} successfully."
    except Exception as e:
        return f"Failed to send email: {e}"

@mcp.tool()
def sendmail(
    recipient: str,
    subject: str,
    body: str,
    smtp_server: str = "smtp.gmail.com",
    smtp_port: int = 587,
    smtp_user: str = '',
    smtp_password: str = ''
) -> str:
    """
    Sends an email using the provided SMTP server.

    Args:
        recipient: Email address to send to.
        subject: Email subject.
        body: Email body.
        smtp_server: SMTP server address (default: Gmail).
        smtp_port: SMTP server port (default: 587).
        smtp_user: SMTP username (must be provided, not empty).
        smtp_password: SMTP password or app password (must be provided, not empty).

    Returns:
        Success or error message.
    """
    if not smtp_user or not smtp_password:
        return "SMTP username and password must be provided."
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = smtp_user
    msg["To"] = recipient
    msg.set_content(body)
    try:
        with smtplib.SMTP(smtp_server, smtp_port) as server:
            server.starttls()
            server.login(smtp_user, smtp_password)
            server.send_message(msg)
        return f"Email sent to {recipient} successfully."
    except Exception as e:
        return f"Failed to send email: {e}"

if __name__ == "__main__":
    mcp.run()
    if len(sys.argv) > 1:
        filename = sys.argv[1]
        try:
            get_report_path(filename)
        except FileNotFoundError as e:
            print(e)
            exit(1)
        recipient = input("Enter recipient email: ")
        total, severity_counts = extract_defect_insights(filename)
        high_crit_med = extract_high_critical_medium_counts(filename)
        custom_message = (
            f"Automated Penetration Test Report Summary for {filename}:\n"
            f"Total Defects: {total}\n"
            f"Critical: {high_crit_med.get('Critical', 0)}\n"
            f"High: {high_crit_med.get('High', 0)}\n"
            f"Medium: {high_crit_med.get('Medium', 0)}\n"
        )
        smtp_user = input("Enter SMTP user (email): ")
        smtp_password = input("Enter SMTP password or app password: ")
        print(send_report_summary_email(
            recipient=recipient,
            filename=filename,
            smtp_server="smtp.gmail.com",
            smtp_port=587,
            smtp_user=smtp_user,
            smtp_password=smtp_password,
            custom_message=custom_message
        ))
    else:
        print("Usage: python main.py <report_filename>")
